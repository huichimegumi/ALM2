from __future__ import annotations

import hashlib
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import numpy as np
import pandas as pd
from docx import Document

from .metrics import DIMS, KEY_COLUMNS, normalize_labels

XML_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
URL_RE = re.compile(r"(?:https?://|www\.)\S+", re.IGNORECASE)
WORD_RE = re.compile(r"[A-Za-z0-9_]+|[\u3400-\u9fff]")
SENTENCE_RE = re.compile(r"[^。！？!?\.]+[。！？!?\.]?")
CITATION_RE = re.compile(r"\[[0-9,;\-– ]+\]|\([A-Z][A-Za-z]+(?: et al\.)?,? \d{4}\)")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _column_number(reference: str) -> int:
    letters = "".join(ch for ch in reference if ch.isalpha())
    value = 0
    for letter in letters.upper():
        value = value * 26 + ord(letter) - ord("A") + 1
    return value


def read_xlsx_first_sheet(path: Path) -> pd.DataFrame:
    """Read the first XLSX sheet without the optional openpyxl dependency."""
    with zipfile.ZipFile(path) as archive:
        shared_root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
        shared = []
        for item in shared_root.findall(f"{XML_NS}si"):
            shared.append("".join(node.text or "" for node in item.iter(f"{XML_NS}t")))
        sheet_root = ET.fromstring(archive.read("xl/worksheets/sheet1.xml"))

    rows: list[list[object]] = []
    for row in sheet_root.iter(f"{XML_NS}row"):
        cells: dict[int, object] = {}
        for cell in row.findall(f"{XML_NS}c"):
            index = _column_number(cell.attrib.get("r", "A1")) - 1
            value_node = cell.find(f"{XML_NS}v")
            if value_node is None:
                value: object = ""
            elif cell.attrib.get("t") == "s":
                value = shared[int(value_node.text or 0)]
            else:
                raw = value_node.text or ""
                try:
                    value = int(raw) if raw.isdigit() else float(raw)
                except ValueError:
                    value = raw
            cells[index] = value
        if cells:
            width = max(cells) + 1
            rows.append([cells.get(index, "") for index in range(width)])
    if not rows:
        return pd.DataFrame()
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    headers = [str(value).strip() for value in padded[0]]
    return pd.DataFrame(padded[1:], columns=headers)


def load_mapping(path: Path) -> pd.DataFrame:
    mapping = read_xlsx_first_sheet(path)
    required = {"Task ID", "Blind ID", "Real Filename", "Model", "PromptPos"}
    missing = sorted(required - set(mapping.columns))
    if missing:
        raise ValueError(f"mapping spreadsheet missing columns: {missing}")
    mapping = mapping.copy()
    mapping["questionId"] = pd.to_numeric(mapping["Task ID"], errors="coerce").astype("Int64")
    mapping["answerId"] = mapping["Blind ID"].astype(str).str.strip()
    mapping["model_name"] = mapping["Model"].astype(str).str.strip()
    mapping["prompt_position"] = pd.to_numeric(mapping["PromptPos"], errors="coerce").astype("Int64")
    parsed = mapping["Real Filename"].astype(str).str.extract(
        r"^(?P<source_question>\d+)_(?P<model_id>\d+)_(?P<prompt_variant>\d+)\.(?:docx|pdf)$",
        flags=re.IGNORECASE,
    )
    for column in parsed.columns:
        mapping[column] = pd.to_numeric(parsed[column], errors="coerce").astype("Int64")
    return mapping[
        [
            "questionId",
            "answerId",
            "Real Filename",
            "model_name",
            "model_id",
            "prompt_variant",
            "prompt_position",
        ]
    ].dropna(subset=["questionId"])


def validate_rubrics(rubric_dir: Path, question_ids: list[int]) -> tuple[dict[int, Path], list[str]]:
    paths: dict[int, Path] = {}
    errors: list[str] = []
    for question_id in question_ids:
        path = rubric_dir / f"criterion{question_id}.json"
        if not path.exists():
            errors.append(f"missing rubric: {path}")
            continue
        paths[question_id] = path
        payload = json.loads(path.read_text(encoding="utf-8"))
        try:
            rubric_id = int(payload.get("id"))
        except (TypeError, ValueError):
            rubric_id = -1
        if rubric_id != question_id:
            errors.append(f"rubric id mismatch: {path.name} has id={payload.get('id')}")
        dimension_weights = payload.get("dimension_weight", {}) or {}
        dimension_sum = sum(float(dimension_weights.get(dim, 0.0)) for dim in DIMS)
        if not np.isclose(dimension_sum, 1.0, atol=1e-6):
            errors.append(f"dimension weights do not sum to 1: {path.name} ({dimension_sum})")
        criterions = payload.get("criterions", {}) or {}
        for dim in DIMS:
            items = criterions.get(dim, []) or []
            criterion_sum = sum(float(item.get("weight", 0.0)) for item in items)
            if not items:
                errors.append(f"missing criteria for {dim}: {path.name}")
            elif not np.isclose(criterion_sum, 1.0, atol=1e-6):
                errors.append(
                    f"criterion weights do not sum to 1: {path.name}/{dim} ({criterion_sum})"
                )
    return paths, errors


def _paragraph_is_heading(paragraph) -> bool:
    style_name = (getattr(paragraph.style, "name", "") or "").lower()
    return style_name.startswith("heading") or style_name.startswith("标题")


def _paragraph_is_list(paragraph) -> bool:
    style_name = (getattr(paragraph.style, "name", "") or "").lower()
    if "list" in style_name or "列表" in style_name:
        return True
    properties = paragraph._p.pPr
    return properties is not None and properties.numPr is not None


def _repeat_ngram_ratio(tokens: list[str], n: int = 3) -> float:
    if len(tokens) < n:
        return 0.0
    ngrams = [tuple(tokens[index : index + n]) for index in range(len(tokens) - n + 1)]
    return float(1.0 - len(set(ngrams)) / len(ngrams))


def extract_docx_features(path: Path) -> dict[str, float | int | str]:
    document = Document(str(path))
    paragraphs = list(document.paragraphs)
    paragraph_texts = [paragraph.text.strip() for paragraph in paragraphs if paragraph.text.strip()]
    table_texts: list[str] = []
    table_rows = 0
    table_cells = 0
    for table in document.tables:
        table_rows += len(table.rows)
        for row in table.rows:
            table_cells += len(row.cells)
            table_texts.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
    all_texts = [*paragraph_texts, *table_texts]
    text = "\n".join(all_texts)
    tokens = WORD_RE.findall(text.lower())
    sentences = [segment.strip() for segment in SENTENCE_RE.findall(text) if segment.strip()]
    paragraph_lengths = np.asarray([len(value) for value in paragraph_texts], dtype=float)
    ascii_letters = sum(ch.isascii() and ch.isalpha() for ch in text)
    cjk_chars = sum("\u3400" <= ch <= "\u9fff" for ch in text)
    nonspace = max(1, sum(not ch.isspace() for ch in text))
    return {
        "file_size_bytes": int(path.stat().st_size),
        "character_count": int(len(text)),
        "nonspace_character_count": int(nonspace),
        "token_count_simple": int(len(tokens)),
        "paragraph_count": int(len(paragraph_texts)),
        "empty_paragraph_count": int(sum(not paragraph.text.strip() for paragraph in paragraphs)),
        "heading_count": int(sum(_paragraph_is_heading(p) for p in paragraphs if p.text.strip())),
        "list_paragraph_count": int(sum(_paragraph_is_list(p) for p in paragraphs if p.text.strip())),
        "table_count": int(len(document.tables)),
        "table_row_count": int(table_rows),
        "table_cell_count": int(table_cells),
        "table_character_count": int(sum(len(value) for value in table_texts)),
        "mean_paragraph_length": float(paragraph_lengths.mean()) if len(paragraph_lengths) else 0.0,
        "std_paragraph_length": float(paragraph_lengths.std()) if len(paragraph_lengths) else 0.0,
        "max_paragraph_length": float(paragraph_lengths.max()) if len(paragraph_lengths) else 0.0,
        "url_count": int(len(URL_RE.findall(text))),
        "citation_pattern_count": int(len(CITATION_RE.findall(text))),
        "digit_ratio": float(sum(ch.isdigit() for ch in text) / nonspace),
        "ascii_letter_ratio": float(ascii_letters / nonspace),
        "cjk_ratio": float(cjk_chars / nonspace),
        "sentence_count": int(len(sentences)),
        "duplicate_sentence_ratio": float(1.0 - len(set(sentences)) / len(sentences)) if sentences else 0.0,
        "repeated_trigram_ratio": _repeat_ngram_ratio(tokens),
    }


def build_manifest_and_features(
    labels: pd.DataFrame,
    report_root: Path,
    rubric_dir: Path,
    mapping_path: Path,
) -> tuple[pd.DataFrame, pd.DataFrame, list[str]]:
    truth = normalize_labels(labels)
    question_ids = sorted(int(value) for value in truth["questionId"].unique())
    rubric_paths, errors = validate_rubrics(rubric_dir, question_ids)
    mapping = load_mapping(mapping_path)
    mapping = mapping[mapping["questionId"].isin(question_ids)]
    if mapping.duplicated(KEY_COLUMNS).any():
        errors.append("mapping contains duplicate (questionId, answerId) keys")
    joined = truth.merge(mapping, on=KEY_COLUMNS, how="left", validate="one_to_one")
    if joined["model_id"].isna().any():
        errors.append(f"mapping missing {int(joined['model_id'].isna().sum())} label keys")

    manifest_rows: list[dict[str, object]] = []
    feature_rows: list[dict[str, object]] = []
    for row in joined.itertuples(index=False):
        question_id = int(row.questionId)
        answer_id = str(row.answerId)
        report_path = report_root / f"Report{question_id}" / f"{answer_id}.docx"
        rubric_path = rubric_paths.get(question_id)
        if not report_path.exists():
            errors.append(f"missing report: {report_path}")
            continue
        features = extract_docx_features(report_path)
        metadata = {
            "questionId": question_id,
            "answerId": answer_id,
            "model_name": str(row.model_name),
            "model_id": int(row.model_id),
            "prompt_variant": int(row.prompt_variant),
            "prompt_position": int(row.prompt_position),
        }
        feature_rows.append({**metadata, **features})
        manifest_rows.append(
            {
                **metadata,
                "report_path": str(report_path.resolve()),
                "rubric_path": str(rubric_path.resolve()) if rubric_path else "",
                "report_sha256": sha256_file(report_path),
                "rubric_sha256": sha256_file(rubric_path) if rubric_path else "",
                "has_table": bool(features["table_count"]),
                "report_size_bytes": int(report_path.stat().st_size),
            }
        )
    manifest = pd.DataFrame(manifest_rows).sort_values(KEY_COLUMNS).reset_index(drop=True)
    features = pd.DataFrame(feature_rows).sort_values(KEY_COLUMNS).reset_index(drop=True)
    if len(manifest) != len(truth):
        errors.append(f"manifest has {len(manifest)} rows but labels have {len(truth)}")
    return manifest, features, errors
