from __future__ import annotations

from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from aeollm_e1.chunking import ChunkingConfig, SourceBlock, chunk_blocks, chunk_docx, count_tokens
from aeollm_e1.dataset import summarize_chunks_by_question

ROOT = Path(__file__).resolve().parents[1]


def test_structural_chunks_keep_headings_separate_and_merge_short_paragraphs() -> None:
    blocks = [
        SourceBlock(0, "heading", "1. Introduction", heading_reason="numbered_heading"),
        SourceBlock(1, "paragraph", "short paragraph"),
        SourceBlock(2, "paragraph", "another short paragraph"),
        SourceBlock(3, "table", "a | b\nc | d", table_rows=["a | b", "c | d"]),
    ]
    chunks = chunk_blocks(blocks, ChunkingConfig(min_chunk_tokens=10, max_chunks=96))
    assert [chunk.type for chunk in chunks] == ["heading", "paragraph", "table"]
    assert chunks[0].source_block_ids == [0]
    assert chunks[1].source_block_ids == [1, 2]
    assert chunks[2].text.startswith("[TABLE]")


def test_long_paragraph_is_split_without_exceeding_token_limit() -> None:
    text = " ".join(f"word{index}" for index in range(1100))
    chunks = chunk_blocks(
        [SourceBlock(0, "paragraph", text)],
        ChunkingConfig(min_chunk_tokens=10, target_chunk_tokens=384, max_chunk_tokens=512),
    )
    assert len(chunks) == 3
    assert max(chunk.token_count for chunk in chunks) <= 512
    assert sum(chunk.token_count for chunk in chunks) == count_tokens(text)


def test_budget_compaction_preserves_all_source_blocks_in_order() -> None:
    blocks = [SourceBlock(index, "heading", f"Heading {index}") for index in range(8)]
    chunks = chunk_blocks(blocks, ChunkingConfig(max_chunks=3))
    assert len(chunks) == 3
    assert [block_id for chunk in chunks for block_id in chunk.source_block_ids] == list(range(8))
    assert any(chunk.overflow_merged for chunk in chunks)


def test_none_chunk_budget_keeps_all_structural_chunks() -> None:
    blocks = [SourceBlock(index, "heading", f"Heading {index}") for index in range(8)]
    chunks = chunk_blocks(blocks, ChunkingConfig(max_chunks=None))
    assert len(chunks) == len(blocks)
    assert not any(chunk.overflow_merged for chunk in chunks)


def test_question_summary_measures_within_question_dispersion() -> None:
    summary = pd.DataFrame(
        {
            "questionId": [1, 1, 2, 2],
            "chunks": [10, 20, 30, 30],
            "total_tokens": [100, 200, 300, 300],
        }
    )
    result = summarize_chunks_by_question(summary).set_index("questionId")
    assert result.loc[1, "chunk_min"] == 10
    assert result.loc[1, "chunk_max_min_ratio"] == 2.0
    assert result.loc[2, "chunk_cv"] == 0.0


def test_real_docx_chunks_have_complete_ordered_provenance() -> None:
    path = ROOT / "data/incoming/google-drive/train/Report1/Doc_001.docx"
    if not path.exists():
        pytest.skip("training DOCX files are not present")
    blocks, chunks = chunk_docx(path)
    assert blocks and chunks
    assert len(chunks) <= 96
    provenance = [block_id for chunk in chunks for block_id in chunk.source_block_ids]
    assert provenance == list(range(len(blocks)))
    assert any(block.type == "heading" for block in blocks)


def test_table_order_is_preserved_in_docx(tmp_path: Path) -> None:
    path = tmp_path / "ordered.docx"
    document = Document()
    document.add_paragraph("Before")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "left"
    table.cell(0, 1).text = "right"
    document.add_paragraph("After")
    document.save(path)
    blocks, _ = chunk_docx(path, ChunkingConfig(min_chunk_tokens=1))
    assert [block.type for block in blocks] == ["paragraph", "table", "paragraph"]
    assert blocks[1].table_rows == ["left | right"]
